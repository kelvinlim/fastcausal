"""
Docx report generation.

Ported from cda_tools2's create_docx_proj.py (PlotnWide class).
"""

import glob
import os
from datetime import datetime
from typing import Optional

from fastcausal.pipeline.config import get_output_dir, get_project_dir


def run_report(
    cfg: dict,
    mode: str = "1wide",
    stub: str = "_label.png",
    verbose: bool = True,
):
    """
    Generate a Word document report from analysis outputs.

    Parameters
    ----------
    cfg : dict
        Loaded config (from load_config).
    mode : str
        Layout mode: "1wide", "2wide", or "3wide".
    stub : str
        Image filename suffix to include (e.g., "_label.png", ".png").
    verbose : bool
        Print progress.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx is required for report generation. "
            "Install with: pip install fastcausal[batch]"
        )

    output_dir = get_output_dir(cfg)
    project_dir = get_project_dir(cfg)
    project_name = cfg.get("GLOBAL", {}).get("name", "project")
    title = cfg.get("GLOBAL", {}).get("title", project_name)
    header = cfg.get("GLOBAL", {}).get("header", "")

    # Get image files
    pattern = os.path.join(output_dir, f"*{stub}")
    files = sorted(glob.glob(pattern))

    if not files:
        if verbose:
            print(f"No files matching '{pattern}'")
        return

    if verbose:
        print(f"Found {len(files)} images matching '*{stub}'")

    # Create document
    doc = Document()

    # Title page
    doc.add_heading(title, level=0)
    if header:
        doc.add_paragraph(header)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Generated: {timestamp}")
    doc.add_paragraph(f"Mode: {mode}")
    doc.add_page_break()

    # Determine images per row
    images_per_row = {"1wide": 1, "2wide": 2, "3wide": 3}.get(mode, 1)
    img_width = {1: Inches(6.0), 2: Inches(3.0), 3: Inches(2.0)}.get(images_per_row, Inches(6.0))

    # Add images
    for i in range(0, len(files), images_per_row):
        batch = files[i:i + images_per_row]

        if images_per_row == 1:
            for f in batch:
                case_name = os.path.basename(f).replace(stub, "")
                doc.add_heading(case_name, level=2)
                try:
                    doc.add_picture(f, width=img_width)
                except Exception as e:
                    doc.add_paragraph(f"Error loading image: {e}")
        else:
            # Multi-column via table
            table = doc.add_table(rows=2, cols=len(batch))
            table.autofit = True

            for j, f in enumerate(batch):
                case_name = os.path.basename(f).replace(stub, "")
                # Header row
                cell = table.cell(0, j)
                cell.text = case_name
                # Image row
                cell = table.cell(1, j)
                try:
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(f, width=img_width)
                except Exception as e:
                    cell.text = f"Error: {e}"

        doc.add_paragraph()  # spacing

    # Appendix: config file
    doc.add_page_break()
    doc.add_heading("Appendix: Configuration", level=1)
    config_path = cfg.get("_config_path", "")
    if config_path and os.path.exists(config_path):
        with open(config_path) as f:
            config_text = f.read()
        doc.add_paragraph(config_text, style="No Spacing")

    # Appendix: prior file
    prior_path = os.path.join(project_dir, "prior.txt")
    if os.path.exists(prior_path):
        doc.add_heading("Appendix: Prior Knowledge", level=1)
        with open(prior_path) as f:
            doc.add_paragraph(f.read(), style="No Spacing")

    # Appendix: case details
    details_path = os.path.join(project_dir, "case_details.csv")
    if os.path.exists(details_path):
        doc.add_heading("Appendix: Case Details", level=1)
        with open(details_path) as f:
            doc.add_paragraph(f.read(), style="No Spacing")

    # Also include effectsize plots
    effectsize_files = sorted(glob.glob(os.path.join(output_dir, "*_effectsize_*.png")))
    if effectsize_files:
        doc.add_page_break()
        doc.add_heading("Effect Size Plots", level=1)
        for f in effectsize_files:
            name = os.path.basename(f)
            doc.add_heading(name, level=3)
            try:
                doc.add_picture(f, width=Inches(6.0))
            except Exception:
                doc.add_paragraph(f"Error loading: {f}")

    # Save document
    doc_filename = f"{project_name}_{mode}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc_path = os.path.join(project_dir, doc_filename)
    doc.save(doc_path)
    if verbose:
        print(f"Report saved: {doc_path}")
